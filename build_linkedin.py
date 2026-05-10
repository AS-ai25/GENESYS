"""
Build GENESYS_LinkedIn_Post.docx — bilingual LinkedIn post (Hebrew + English)
"""
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUT = r"C:\Users\amirs\Documents\ClaudCode\Genesys\GENESYS_LinkedIn_Post.docx"
doc = Document()

sec = doc.sections[0]
sec.page_width    = Cm(21.0);  sec.page_height   = Cm(29.7)
sec.left_margin   = Cm(3.0);   sec.right_margin  = Cm(3.0)
sec.top_margin    = Cm(2.5);   sec.bottom_margin = Cm(2.5)

NAVY  = RGBColor(0x0a, 0x29, 0x66)
BLUE  = RGBColor(0x00, 0x77, 0xb5)   # LinkedIn blue
DARK  = RGBColor(0x1a, 0x1a, 0x1a)
GREY  = RGBColor(0x55, 0x55, 0x55)
WHITE = RGBColor(0xff, 0xff, 0xff)

def set_rtl(p):
    pPr = p._p.get_or_add_pPr()
    b = OxmlElement('w:bidi'); b.set(qn('w:val'), '1'); pPr.append(b)
    j = OxmlElement('w:jc');   j.set(qn('w:val'), 'right'); pPr.append(j)

def set_run_rtl(run):
    rPr = run._r.get_or_add_rPr()
    r = OxmlElement('w:rtl'); r.set(qn('w:val'), '1'); rPr.append(r)

def shd(p, fill):
    pPr = p._p.get_or_add_pPr()
    s = OxmlElement('w:shd')
    s.set(qn('w:val'), 'clear'); s.set(qn('w:color'), 'auto'); s.set(qn('w:fill'), fill)
    pPr.append(s)

def divider(color='0077b5'):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8); p.paragraph_format.space_after = Pt(8)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bot = OxmlElement('w:bottom')
    bot.set(qn('w:val'), 'single'); bot.set(qn('w:sz'), '4')
    bot.set(qn('w:space'), '1');    bot.set(qn('w:color'), color)
    pBdr.append(bot); pPr.append(pBdr)

def lp(text, size=11, bold=False, italic=False, color=None,
        align=WD_ALIGN_PARAGRAPH.LEFT, sb=0, sa=5, rtl=False, indent=0):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_before = Pt(sb)
    p.paragraph_format.space_after  = Pt(sa)
    if indent: p.paragraph_format.left_indent = Cm(indent)
    if rtl: set_rtl(p)
    run = p.add_run(text)
    run.font.size = Pt(size); run.font.bold = bold; run.font.italic = italic
    run.font.color.rgb = RGBColor(*color) if color else DARK
    if rtl: set_run_rtl(run)
    return p

# ─────────────────────────────────────────────────────────────
#  HEADER BANNER
# ─────────────────────────────────────────────────────────────
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(0); p.paragraph_format.space_after = Pt(0)
p.paragraph_format.left_indent = Cm(-0.5); p.paragraph_format.right_indent = Cm(-0.5)
shd(p, '0077b5')
run = p.add_run("  🧬  GENESYS  ·  LinkedIn Announcement  🧬  ")
run.font.size = Pt(16); run.font.bold = True; run.font.color.rgb = WHITE

doc.add_paragraph().paragraph_format.space_after = Pt(4)

# ─────────────────────────────────────────────────────────────
#  HEBREW VERSION
# ─────────────────────────────────────────────────────────────
p = doc.add_paragraph()
set_rtl(p)
p.paragraph_format.space_before = Pt(4); p.paragraph_format.space_after = Pt(2)
run = p.add_run("🇮🇱  גרסה עברית  —  LinkedIn Post")
run.font.size = Pt(10); run.font.bold = True; run.font.color.rgb = BLUE
set_run_rtl(run)

# Hebrew post box
p = doc.add_paragraph()
set_rtl(p)
p.paragraph_format.space_before = Pt(4); p.paragraph_format.space_after = Pt(0)
p.paragraph_format.left_indent = Cm(0.4); p.paragraph_format.right_indent = Cm(0.4)
shd(p, 'EBF5FB')
run = p.add_run("🚀 מה קורה כשנותנים לכל גן בתא חיה זהות משלו, זיכרון, ואישיות — ואז שולחים אותו לעבוד עם 4,761 עמיתים?")
run.font.size = Pt(13); run.font.bold = True; run.font.color.rgb = NAVY
set_run_rtl(run)

he_lines = [
    "בניתי את GENESYS — סימולטור תא חי של E. coli K-12 עם 4,762 סוכני גן אינדיבידואליים.",
    "",
    "כל גן הוא סוכן AI עצמאי:",
    "   🔬 זהות מולקולרית מלאה — פונקציה, משקל, מיקום תת-תאי",
    "   🧠 תגובה מבוססת LLM (Claude / llama3) — לכל שינוי סביבתי",
    "   🔗 מחובר ל-24,000 קשרי חלבון-חלבון (PPI) ידועים",
    "   💾 זיכרון של 10 תגובות אחרונות",
    "",
    "משנים את הטמפרטורה ל-95°C? → גנים של תגובת חום מואצים מיידית.",
    "מוסיפים ציפרופלוקסצין? → רשת SOS שלמה מופעלת בשניות.",
    "מדלדלים את כל חומרי המזון? → 220 גנים נכנסים למצב הישרדות.",
    "",
    "🌿 הרעיון המרכזי: LLM-agents לא שייכים רק לעולם האנושי.",
    "הם יכולים לייצג כל ישות ביולוגית — גן, תא, אורגניזם, מערכת אקולוגית.",
    "E. coli היא רק ההתחלה. Arabidopsis, שמר, מיקרוביום מעי — הכל אפשרי.",
    "",
    "📊 מה שבניתי:",
    "   • 4,762 גנים כסוכנים — בזמן אמת, ב-3D",
    "   • ויזואליזציה של רשת PPI + אופרונים",
    "   • גודל כדור = MW × Protein Abundance (PPM) — חשיבות תאית בצבע וצורה",
    "   • תמיכה ב-Ollama llama3 מקומי — ללא API, ללא עלות",
    "   • AI environment interpreter — כותבים 'תרחיש בטבע' ו-LLM מגדיר את הסביבה",
    "",
    "זה לא רק כלי — זה הוכחת עקרון שמשנה את הדרך שבה אנחנו חושבים על ביולוגיה חישובית.",
    "",
    "💬 מעוניינ/ת לשתף פעולה? לשמוע עוד? לנסות את הקוד?",
    "השאירו תגובה או שלחו הודעה 👇",
    "",
    "#AI #SystemsBiology #LLM #Bioinformatics #ArtificialIntelligence #Ecoli",
    "#ComputationalBiology #AgentBasedModeling #OpenSource #BioTech #Innovation",
]

for line in he_lines:
    p = doc.add_paragraph()
    set_rtl(p)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(1 if line else 4)
    p.paragraph_format.left_indent  = Cm(0.4)
    p.paragraph_format.right_indent = Cm(0.4)
    shd(p, 'EBF5FB')
    run = p.add_run(line)
    is_hashtag = line.startswith('#')
    is_header  = line.startswith('🌿') or line.startswith('📊') or line.startswith('🚀')
    is_bullet  = line.strip().startswith('•') or line.strip().startswith('🔬') or \
                 line.strip().startswith('🧠') or line.strip().startswith('🔗') or \
                 line.strip().startswith('💾')
    run.font.size = Pt(9 if is_hashtag else 10)
    run.font.bold = is_header
    run.font.color.rgb = BLUE if is_hashtag else (NAVY if is_header else DARK)
    set_run_rtl(run)

divider()

# ─────────────────────────────────────────────────────────────
#  ENGLISH VERSION
# ─────────────────────────────────────────────────────────────
lp("🌐  English Version  —  LinkedIn Post",
   size=10, bold=True, color=(0x00, 0x77, 0xb5), sb=4, sa=2)

# English post box
p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(4); p.paragraph_format.space_after = Pt(0)
p.paragraph_format.left_indent  = Cm(0.4); p.paragraph_format.right_indent = Cm(0.4)
shd(p, 'EBF5FB')
run = p.add_run(
    "🚀 What happens when you give every gene in a living cell its own identity, memory, "
    "and personality — and set all 4,762 of them to work together in real time?"
)
run.font.size = Pt(13); run.font.bold = True; run.font.color.rgb = NAVY

en_lines = [
    "I built GENESYS — a living-cell simulator of E. coli K-12 powered by 4,762 individual gene agents.",
    "",
    "Each gene is an autonomous AI agent with:",
    "   🔬 Full molecular identity — function, weight, subcellular location",
    "   🧠 LLM-driven responses (Claude / llama3) to every environmental change",
    "   🔗 Connections to 24,000 experimentally validated protein–protein interactions",
    "   💾 A rolling memory of its last 10 reactions",
    "",
    "Raise the temperature to 95 °C? → Heat-shock genes spike within seconds.",
    "Add ciprofloxacin? → The entire SOS network lights up across the 3D graph.",
    "Deplete all nutrients? → 220 genes enter survival mode simultaneously.",
    "",
    "🌿 The core idea: LLM-agents aren't just for human tasks.",
    "Any biological entity — gene, cell, organism, ecosystem — can be an agent.",
    "E. coli is the proof of concept. Arabidopsis, yeast, the gut microbiome — all within reach.",
    "",
    "📊 What I built:",
    "   • 4,762 gene agents running live in an interactive 3D cell",
    "   • Real-time PPI network + operon co-regulation visualisation",
    "   • Node size = MW × Protein Abundance PPM (cellular importance made visible)",
    "   • Full Ollama llama3 support — local inference, no API key, no cost",
    "   • AI environment interpreter — describe a natural scenario in plain text,",
    "     and the LLM sets every environmental parameter automatically",
    "",
    "This isn't just a tool — it's a proof of principle that reframes how we think",
    "about computational biology and the reach of language model agents.",
    "",
    "I'm publishing the findings as a scientific paper and opening the code.",
    "If you work in systems biology, bioinformatics, or AI — I'd love to connect.",
    "",
    "💬 Interested in collaborating, testing the code, or just talking science?",
    "Drop a comment or send a message 👇",
    "",
    "#AI #SystemsBiology #LLM #Bioinformatics #ArtificialIntelligence #Ecoli",
    "#ComputationalBiology #AgentBasedModeling #OpenSource #BioTech #Innovation",
]

for line in en_lines:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(1 if line else 4)
    p.paragraph_format.left_indent  = Cm(0.4)
    p.paragraph_format.right_indent = Cm(0.4)
    shd(p, 'EBF5FB')
    run = p.add_run(line)
    is_hashtag = line.startswith('#')
    is_header  = any(line.startswith(e) for e in ['🌿', '📊', '🚀'])
    run.font.size = Pt(9 if is_hashtag else 10)
    run.font.bold = is_header
    run.font.color.rgb = BLUE if is_hashtag else (NAVY if is_header else DARK)

divider()

# ─────────────────────────────────────────────────────────────
#  COPY-PASTE TIPS
# ─────────────────────────────────────────────────────────────
p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(4); p.paragraph_format.space_after = Pt(2)
run = p.add_run("📋  Copy-Paste Tips for LinkedIn")
run.font.size = Pt(10); run.font.bold = True; run.font.color.rgb = GREY

tips = [
    "Post the Hebrew version for Israeli network, English for international reach — or post both in one post (Hebrew first, then a divider ———, then English).",
    "Add 2–3 screenshots of the running app: the 3D cell graph with active gene connections, a zoomed-in cluster with gene labels, and the feed showing LLM reactions.",
    "Tag relevant hashtags in the first comment rather than the post body to keep the text clean.",
    "Best posting time: Tuesday–Thursday, 8–10 AM or 12–1 PM local time.",
    "The hook line (first sentence) is the most critical — LinkedIn shows only 2–3 lines before 'see more'. Both versions open with a question to maximise click-through.",
]
for tip in tips:
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run(tip)
    run.font.size = Pt(9); run.font.color.rgb = GREY

doc.save(OUT)
print(f"Saved: {OUT}")
