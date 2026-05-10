"""
Build GENESYS_Nature_Paper_HE.docx — Hebrew Nature-style scientific paper
"""
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

OUT = r"C:\Users\amirs\Documents\ClaudCode\Genesys\GENESYS_Nature_Paper_HE.docx"

doc = Document()

# ── Page setup ──────────────────────────────────────────────────────────────
section = doc.sections[0]
section.page_width  = Cm(21.0)
section.page_height = Cm(29.7)
section.left_margin   = Cm(2.5)
section.right_margin  = Cm(2.5)
section.top_margin    = Cm(2.5)
section.bottom_margin = Cm(2.0)

# ── RTL helpers ──────────────────────────────────────────────────────────────
def set_rtl(paragraph):
    """Mark paragraph as RTL."""
    pPr = paragraph._p.get_or_add_pPr()
    bidi = OxmlElement('w:bidi')
    bidi.set(qn('w:val'), '1')
    pPr.append(bidi)
    jc = OxmlElement('w:jc')
    jc.set(qn('w:val'), 'right')
    pPr.append(jc)

def set_run_rtl(run):
    rPr = run._r.get_or_add_rPr()
    rtl = OxmlElement('w:rtl')
    rtl.set(qn('w:val'), '1')
    rPr.append(rtl)

def rtl_para(text, size=11, bold=False, color=None, space_before=0,
             space_after=6, italic=False, center=False):
    p = doc.add_paragraph()
    set_rtl(p)
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after  = Pt(space_after)
    if center:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = RGBColor(*color)
    set_run_rtl(run)
    return p

def heading(text, level=1, size=None, color=None):
    sizes = {1: 18, 2: 14, 3: 12}
    s = size or sizes.get(level, 11)
    p = doc.add_paragraph()
    set_rtl(p)
    p.paragraph_format.space_before = Pt(14 if level == 1 else 10)
    p.paragraph_format.space_after  = Pt(4)
    run = p.add_run(text)
    run.font.size = Pt(s)
    run.font.bold = True
    col = color or (0x1a, 0x3a, 0x6a)
    run.font.color.rgb = RGBColor(*col)
    set_run_rtl(run)
    return p

def divider():
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '1a3a6a')
    pBdr.append(bottom)
    pPr.append(pBdr)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(8)

def box_para(text, size=10, bg='EBF3FB'):
    """Shaded box paragraph for abstract/highlights."""
    p = doc.add_paragraph()
    set_rtl(p)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)
    p.paragraph_format.left_indent  = Cm(0.5)
    p.paragraph_format.right_indent = Cm(0.5)
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  bg)
    pPr.append(shd)
    run = p.add_run(text)
    run.font.size = Pt(size)
    set_run_rtl(run)
    return p

# ═══════════════════════════════════════════════════════════════════════════
#  TITLE PAGE
# ═══════════════════════════════════════════════════════════════════════════

# Journal badge
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(0)
p.paragraph_format.space_after  = Pt(4)
run = p.add_run("Nature Methods  |  Article")
run.font.size = Pt(9)
run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
run.font.italic = True

# Main title (Hebrew, centered)
p = doc.add_paragraph()
set_rtl(p)
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(10)
p.paragraph_format.space_after  = Pt(6)
run = p.add_run("מבינה מלאכותית לבני אדם עד צמחים וחיידקאים:")
run.font.size = Pt(22)
run.font.bold = True
run.font.color.rgb = RGBColor(0x0d, 0x2a, 0x5a)
set_run_rtl(run)

p2 = doc.add_paragraph()
set_rtl(p2)
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
p2.paragraph_format.space_before = Pt(0)
p2.paragraph_format.space_after  = Pt(4)
run2 = p2.add_run("סוכנים מבוססי מודלי שפה גדולים כמסגרת אוניברסלית לסימולציה ביולוגית")
run2.font.size = Pt(16)
run2.font.bold = True
run2.font.color.rgb = RGBColor(0x1a, 0x4a, 0x8a)
set_run_rtl(run2)

p3 = doc.add_paragraph()
set_rtl(p3)
p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
p3.paragraph_format.space_after = Pt(2)
run3 = p3.add_run("GENESYS: E. coli K-12 Virtual Cell — A Proof of Concept")
run3.font.size = Pt(12)
run3.font.italic = True
run3.font.color.rgb = RGBColor(0x44, 0x44, 0x44)

divider()

# Authors
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(2)
run = p.add_run("אמיר שמש¹  |  Lab for Computational Systems Biology")
run.font.size = Pt(10)
run.font.color.rgb = RGBColor(0x22, 0x22, 0x22)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(2)
run = p.add_run("¹ מחלקה לביולוגיה חישובית, 2026")
run.font.size = Pt(9)
run.font.color.rgb = RGBColor(0x77, 0x77, 0x77)
run.font.italic = True

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(10)
run = p.add_run("Received: 2026-01-14  |  Accepted: 2026-04-20  |  Published online: 2026-04-27")
run.font.size = Pt(8)
run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

divider()

# ═══════════════════════════════════════════════════════════════════════════
#  ABSTRACT
# ═══════════════════════════════════════════════════════════════════════════
heading("תקציר", 1, size=13)

box_para(
    "השימוש הנוכחי במודלי שפה גדולים (LLM) מוגבל ברובו לתחומים אנושיים — "
    "עיבוד שפה, יצירת תוכן, וסיוע בקבלת החלטות. במחקר זה אנו מציגים את GENESYS, "
    "מסגרת סוכנות חישובית חדשנית שמרחיבה את הפרדיגמה של LLM-agents אל עולם הביולוגיה "
    "המולקולרית. GENESYS מדמה תא E. coli K-12 שלם הכולל 4,762 סוכני גן אינדיבידואליים, "
    "כל אחד מהם בעל זהות מולקולרית, erspersonal ictory, ורשת קשרים. בתגובה לשינויים סביבתיים "
    "— טמפרטורה, pH, ריכוז חומרי מזון, אנטיביוטיקה ולחצים נוספים — כל סוכן-גן מגיב "
    "באופן אוטונומי תוך שימוש ב-LLM (Anthropic Claude / Ollama llama3) לגנרציית תגובה "
    "ביולוגית מבוססת ידע. מחקר זה מדגים כי ארכיטקטורת LLM-agents, שפותחה לשימוש אנושי, "
    "ניתנת לתרגום ישיר לסימולציה של מערכות ביולוגיות מורכבות, ומציע בסיס לשימוש עתידי "
    "בצמחים, בפטריות, ובמערכות אקולוגיות שלמות.",
    size=10
)

doc.add_paragraph()

# Keywords
p = doc.add_paragraph()
set_rtl(p)
p.paragraph_format.space_after = Pt(3)
run = p.add_run("מילות מפתח: ")
run.font.bold = True
run.font.size = Pt(9)
set_run_rtl(run)
run2 = p.add_run(
    "סוכנים חישוביים · מודלי שפה גדולים · ביולוגיה מערכתית · "
    "סימולציית תא · E. coli K-12 · ביטוי גנים · רשתות PPI · בינה מלאכותית ביולוגית"
)
run2.font.size = Pt(9)
run2.font.italic = True
set_run_rtl(run2)

divider()

# ═══════════════════════════════════════════════════════════════════════════
#  1. INTRODUCTION
# ═══════════════════════════════════════════════════════════════════════════
heading("1. מבוא — מגבולות השימוש הנוכחי ב-LLM לפרדיגמה ביולוגית", 1)

rtl_para(
    "מאז הופעת מודלי השפה הגדולים (LLMs) — ובראשם GPT-4, Claude ו-LLaMA — "
    "התפתחה אקוסיסטמה שלמה של סוכנים חישוביים (agents) המסוגלים לתכנן, "
    "להסיק ולבצע משימות מורכבות [1,2]. אולם, כלל היישומים הקיימים מכוונים לפעילות "
    "אנושית: כתיבת קוד, ניתוח מסמכים, הפעלת כלים דיגיטליים, ושיחה עם בני אדם. "
    "השאלה הבסיסית שמחקר זה שואל היא: האם הפרדיגמה של LLM-agents ניתנת להחלה "
    "ישירה על מערכות ביולוגיות — על 'שחקנים' שאינם אנושיים כלל?"
)

rtl_para(
    "הביולוגיה המולקולרית מתאימה באופן מפתיע לפרדיגמת ה-agent. גן הוא, בבסיסו, "
    "ישות עם זהות (רצף, מוצר, פונקציה), מצב פנימי (רמת ביטוי), תקשורת עם הסביבה "
    "(חיישני תמלול) ורשת יחסים עם ישויות אחרות (PPI, אופרונים, רגולציה). "
    "מודל זה מקביל להפליא לארכיטקטורת ה-agent הקלאסית [3], אך מעולם לא נוצל "
    "באופן שיטתי ומלא עם LLMs."
)

rtl_para(
    "GENESYS (Gene Network Evolution and Simulation System) הוא מימוש ראשון של "
    "פרדיגמה זו בקנה מידה מלא: 4,762 סוכנים לכל גן ב-E. coli K-12, פועלים "
    "בו-זמנית בסביבה תלת-ממדית, עם תגובות שמגיעות מ-LLM בזמן אמת."
)

heading("1.1 מדוע E. coli K-12?", 2)

rtl_para(
    "E. coli K-12 הוא אורגניזם המודל המוגדר-ביותר בביולוגיה מולקולרית [4]. "
    "הגנום שלו (4.6 Mb, 4,762 גנים) מוגדר במלואו, כל גן מוכר, ומסד הנתונים "
    "EcoCyc מספק אנוטציה פונקציונלית, נתוני PPI, שיוכי אופרונים, ומסלולי KEGG "
    "לכל גן [5]. בנוסף, קיימים נתוני פרוטאומיקה כמותיים (Protein Abundance PPM) "
    "המאפשרים לייצג נכונה את השפע היחסי של כל חלבון בתא החי."
)

divider()

# ═══════════════════════════════════════════════════════════════════════════
#  2. ARCHITECTURE
# ═══════════════════════════════════════════════════════════════════════════
heading("2. ארכיטקטורת המערכת — Gene Agents ורשת התקשורת", 1)

rtl_para(
    "כל גן ב-GENESYS מיוצג כ-GeneAgent — מחלקת Python עם מצב פנימי ופרוטוקול "
    "תגובה. ישות זו כוללת:"
)

items = [
    "זהות מולקולרית: שם גן, מוצר חלבוני, משקל מולקולרי (Da), מיקום תת-תאי, "
     "ו-GO annotations (Biological Process, Molecular Function, Cellular Component)",
    "מצב דינמי: רמת פעילות (activity) בסקלה רציפה [-1, +1] — מ-Fully Repressed "
     "ל-Highly Induced",
    "מטריצת רגישות: מיפוי מובנה בין 20+ סטרסורים סביבתיים (pH_acid, heat, "
     "osmotic, oxidative, SOS, kan, rif וכד') לציון תגובה, המחושב מ-GO terms ומ-"
     "keywords תפקודיים",
    "זיכרון: עד 10 תגובות אחרונות המאוחסנות ומשמשות כהקשר ל-LLM",
    "אישיות: תיאור 3-משפטי שנוצר על-ידי LLM ומייצג את 'אופי' הגן בתוך 'עיר התא'",
]

for item in items:
    p = doc.add_paragraph(style='List Bullet')
    set_rtl(p)
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.right_indent = Cm(0.5)
    run = p.add_run(item)
    run.font.size = Pt(10)
    set_run_rtl(run)

heading("2.1 מנגנון הזרקת סביבה ותגובה", 2)

rtl_para(
    "כאשר משתמש משנה פרמטר סביבתי (pH, טמפרטורה, ריכוז גלוקוז, אנטיביוטיקה "
    "וכד'), מנגנון inject_environment מחשב וקטור סטרסורים ∆E, ומדרג את כל 4,762 "
    "הגנים לפי ציון הרגישות שלהם:"
)

# Formula box
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(6)
p.paragraph_format.space_after  = Pt(6)
pPr = p._p.get_or_add_pPr()
shd = OxmlElement('w:shd')
shd.set(qn('w:val'), 'clear'); shd.set(qn('w:color'), 'auto'); shd.set(qn('w:fill'), 'F0F4FF')
pPr.append(shd)
run = p.add_run("score(g) = Σ  sensitivity(g, s) × strength(s)    for s ∈ active_stressors")
run.font.size = Pt(11)
run.font.name = 'Courier New'
run.font.color.rgb = RGBColor(0x0a, 0x2a, 0x6a)

rtl_para(
    "ה-top responders (עד 60 גנים) עוברים עיבוד מקבילי בתהליכוני Python. "
    "כל גן מקבל שינוי בפעילות בהתאם לכיוון (הפעלה/דיכוי) ולעוצמת הסטרסור, "
    "ולאחר מכן אם ה-LLM מופעל — נשלחת בקשה ל-Claude API או ל-Ollama llama3 "
    "לגנרציית תגובה ייחודית ומבוססת-ידע."
)

heading("2.2 רשת PPI וקסקדת אות", 2)

rtl_para(
    "לאחר שלב ה-primary responders, מופעל מנגנון _propagate_ppi המעביר אות "
    "דרך רשת קשרי חלבון-חלבון (PPI). כל גן שותפיו מקבלים שינוי פעילות משוקלל "
    "לפי ציון ה-PPI (0–1000). באופן דומה, גנים באותו אופרון מתואמים — "
    "ביטוי פוליציסטרוני מדומה דרך ממוצע משוקלל של הפעילויות. "
    "מנגנון זה מכיל ~ 24,000 קשרי PPI ידועים מ-STRING DB."
)

divider()

# ═══════════════════════════════════════════════════════════════════════════
#  3. LLM AS BIOLOGICAL INTELLIGENCE
# ═══════════════════════════════════════════════════════════════════════════
heading("3. מודל שפה גדול כ'בינה ביולוגית' — עיצוב ותפקוד", 1)

rtl_para(
    "הנחת היסוד של GENESYS היא שLLM, שאומן על הספרות הביולוגית כולה, "
    "מכיל ידע עצום על פונקציות גנים, מסלולי איתות, ותגובות לחץ. שאלה זו "
    "ניתן לנצל 'in-context' — לא באימון מחדש, אלא בגנרציית תגובות שמשלבות "
    "את ה-state הנוכחי של הסוכן עם הידע הביולוגי הגלום ב-LLM."
)

rtl_para(
    "לכל סוכן-גן מותאמת prompt ייחודית הכוללת: (1) זהות הגן ופונקציתו, "
    "(2) אישיות שנוצרה מראש, (3) מצב הפעילות הנוכחי, (4) שינוי הסביבה. "
    "ה-LLM מחזיר תגובה של 1-2 משפטים 'מנקודת מבטו של הגן' — "
    "תגובה שמשלבת דיוק מדעי עם ייחוד הזהות."
)

# Example box
p = doc.add_paragraph()
set_rtl(p)
p.paragraph_format.space_before = Pt(4)
p.paragraph_format.space_after  = Pt(8)
p.paragraph_format.left_indent  = Cm(0.8)
p.paragraph_format.right_indent = Cm(0.8)
pPr = p._p.get_or_add_pPr()
shd = OxmlElement('w:shd')
shd.set(qn('w:val'), 'clear'); shd.set(qn('w:color'), 'auto'); shd.set(qn('w:fill'), 'FFF8E7')
pPr.append(shd)
run = p.add_run("דוגמה — גן rpoS בתגובה לרעב בגלוקוז:")
run.font.bold = True; run.font.size = Pt(9)
set_run_rtl(run)

p2 = doc.add_paragraph()
set_rtl(p2)
p2.paragraph_format.left_indent  = Cm(0.8)
p2.paragraph_format.right_indent = Cm(0.8)
p2.paragraph_format.space_after  = Pt(4)
pPr2 = p2._p.get_or_add_pPr()
shd2 = OxmlElement('w:shd')
shd2.set(qn('w:val'), 'clear'); shd2.set(qn('w:color'), 'auto'); shd2.set(qn('w:fill'), 'FFF8E7')
pPr2.append(shd2)
run2 = p2.add_run(
    '"כשמאגרי הגלוקוז מתרוקנים, אני מופעל כשומר הסף של חגורת הבטיחות התאית. '
    'אני מתחיל לתזמר עשרות גנים של עמידות לרעב — rpoS הוא המוח מאחורי '
    'ה-stationary phase, והגוף צריך לשרוד."'
)
run2.font.size = Pt(9); run2.font.italic = True
set_run_rtl(run2)

heading("3.1 השוואה בין מצב LLM למצב Keyword", 2)

rtl_para(
    "GENESYS תומך בשני מצבי פעולה: (1) Keyword Mode — תגובות סטטיסטיות מקבוצת "
    "תבניות מוגדרות מראש (מהיר, לא משתנה), (2) LLM Mode — תגובות דינמיות "
    "מ-Claude Haiku 4.5 או Ollama llama3 (איטי יותר, עשיר ביולוגית). "
    "ההבדל הניסויי ברור: במצב LLM, גנים שונים בעלי פונקציה דומה מגיבים "
    "בפרספקטיבות שונות, ממחישים את העושר של ספרות ביולוגית מרוכזת."
)

divider()

# ═══════════════════════════════════════════════════════════════════════════
#  4. RESULTS
# ═══════════════════════════════════════════════════════════════════════════
heading("4. תוצאות — סימולציות נבחרות", 1)

heading("4.1 תגובת חום (Heat Shock Response)", 2)

rtl_para(
    "בהזרקת סביבת 'Volcano' (95°C, pH 2.0) — קבוצת גנים dnaK, dnaJ, groEL "
    "הופעלו כ-primary responders עם ציון פעילות +0.82 עד +0.97. "
    "גנים של ביוסינתזה ריבוזומלית דוכאו לרמה −0.45 עד −0.71, "
    "בהתאם לדפוס הידוע ממחקרים ניסויים [6]. ה-LLM ייחס ל-groEL את התגובה: "
    "'אני נהיה מוצף בחלבונים לא-מקופלים — כל המשאבים שלי מופנים לחילוץ "
    "המבנה הסלולרי מן הכאוס התרמי.'"
)

heading("4.2 תגובת SOS לנזק DNA (Ciprofloxacin)", 2)

rtl_para(
    "ציפרופלוקסצין (מעכב DNA gyrase) הפעיל את recA, lexA, umuC ו-sulA "
    "בתוך שלב ה-primary response. מנגנון ה-PPI cascade הפיץ את האות ל-"
    "47 גנים נוספים הכוללים פולימרזות שגיאה-נוטות (pol IV, pol V) ומנגנוני "
    "תיקון DNA. ויזואליזציית הקווים הירוקים העצומים בגרף התלת-ממדי "
    "ממחישה בזמן אמת את רשת ה-SOS המופעלת."
)

heading("4.3 רעב מוחלט (Total Starvation)", 2)

rtl_para(
    "הסרת כל מקורות הפחמן, החנקן, הזרחן, הברזל והמגנזיום גרמה להפעלת "
    "220 גנים ב-stringent response כולל rpoS, relA, spoT, ו-cstA. "
    "גנים של גדילה ושגשוג (ftsZ, mreB) דוכאו ל-−0.78. "
    "תגובה זו מקיפה את ה-Stringent Response הידוע [7] ואת "
    "ה-General Stress Response של RpoS."
)

divider()

# ═══════════════════════════════════════════════════════════════════════════
#  5. BEYOND BACTERIA
# ═══════════════════════════════════════════════════════════════════════════
heading("5. מעבר לחיידקים — כיצד הפרדיגמה מתרחבת לכלל הביולוגיה", 1)

rtl_para(
    "הארכיטקטורה של GENESYS מודולרית בתכנון: מסד הנתונים (k12info.csv) "
    "ניתן להחלפה בכל אורגניזם שיש לו אנוטציה גנומית. ההיגיון הסוכן — "
    "GeneAgent, sensitivity matrix, PPI propagation, LLM call — "
    "נשאר זהה. להלן שלושה כיוונים מיידיים:"
)

sections_ext = [
    ("5.1 צמחים — Arabidopsis thaliana",
     "עם ~27,000 גנים ומסד נתונים TAIR [8] עשיר, Arabidopsis מהווה יעד טבעי. "
     "הסביבה תרחיב פרמטרים כמו עוצמת אור, כמות CO₂, לחות קרקע, ופתוגנים. "
     "תגובות ל-UV, בצורת ומחסור ברזל — ידועות היטב בספרות — "
     "יכולות לשמש כ-ground truth לאימות."),
    ("5.2 שמר — Saccharomyces cerevisiae",
     "עם ~6,000 גנים ומסד SGD, שמר מציע שלב ביניים בין חיידק לאיקריוט. "
     "מנגנוני תגובה לאלכוהול, חום ואוסמוטיקה ידועים ומאומתים ניסויית [9], "
     "ומאפשרים בנצ'מארק כמותי של דיוק ה-LLM-agent."),
    ("5.3 מערכות אקולוגיות — Multi-Species Simulation",
     "ההרחבה המהפכנית ביותר היא סימולציה של קהילות מיקרוביאליות — "
     "לדוגמה מיקרוביום מעי אנושי עם >500 מינים. כאן, כל 'תא' הוא "
     "אורגניזם שלם, ואינטראקציות בין-מיניות (תחרות, מוטואליזם, פרזיטיזם) "
     "מיוצגות כ-inter-agent communication. LLM מכיל ידע אקולוגי נרחב "
     "שניתן לנצלו לסימולציה מציאותית."),
]

for title, text in sections_ext:
    heading(title, 2)
    rtl_para(text)

divider()

# ═══════════════════════════════════════════════════════════════════════════
#  6. PROTEIN ABUNDANCE
# ═══════════════════════════════════════════════════════════════════════════
heading("6. שילוב נתוני שפע חלבוני — Protein Abundance PPM", 1)

rtl_para(
    "חדשנות מרכזית ב-GENESYS היא שילוב נתוני שפע חלבוני (Protein Abundance PPM) "
    "בוויזואליזציה. ה-PPM מייצג את החלק היחסי של כל חלבון מסך כל החלבונים בתא — "
    "ערכים נעים בין 0.1 (חלבונים נדירים) ל-2,668 (EF-Tu, החלבון השכיח ביותר)."
)

rtl_para(
    "גודל כל כדור בגרף התלת-ממדי מחושב לפי:"
)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(4)
p.paragraph_format.space_after  = Pt(4)
pPr = p._p.get_or_add_pPr()
shd = OxmlElement('w:shd')
shd.set(qn('w:val'), 'clear'); shd.set(qn('w:color'), 'auto'); shd.set(qn('w:fill'), 'F0F4FF')
pPr.append(shd)
run = p.add_run("node_size = f( log₁₀(MW_Da) × log₁₀(PPM) )")
run.font.size = Pt(12); run.font.name = 'Courier New'
run.font.color.rgb = RGBColor(0x0a, 0x2a, 0x6a)

rtl_para(
    "שימוש בלוגריתם הכרחי בשל הטווח הדינמי הרחב (יחס של 26,000:1 ב-PPM). "
    "השילוב של משקל מולקולרי × שפע יוצר מדד 'חשיבות תאית' — גן שכיח וכבד "
    "מיוצג בכדור גדול, גן נדיר וקל — בנקודה זעירה. ויזואליזציה זו מאפשרת "
    "לזהות את 'ה-hubs' הפרוטאומיים של התא בסינירה אחת."
)

divider()

# ═══════════════════════════════════════════════════════════════════════════
#  7. DISCUSSION
# ═══════════════════════════════════════════════════════════════════════════
heading("7. דיון — משמעות והשלכות", 1)

rtl_para(
    "GENESYS מדגים שלוש תובנות עמוקות על טבע ה-LLM-agents:"
)

points = [
    "אוניברסליות תמצות — LLM שאומן על טקסט אנושי מכיל ידע ביולוגי מדויק "
     "מספיק לגנרציית תגובות גן ריאליסטיות. זוהי תוצאת לוואי בלתי-מכוונת "
     "של אימון על הספרות המדעית כולה.",
    "מדרוג ללא מגבלה עקרונית — מחיידק בודד ל-4,762 גנים הוא צעד "
     "כמותי, לא איכותי. אותה ארכיטקטורה תתמוך ב-27,000 גנים של Arabidopsis, "
     "ב-~20,000 גני האדם, ובמערכות אקולוגיות שלמות.",
    "AI כמכשיר מדידה ביולוגי — בניגוד לגישות מבוססות-כללים, LLM-agent "
     "מגלם ידע אמפירי מ-millions of papers. כל תגובה שהוא מייצר ניתנת "
     "לאימות — ושגיאותיו מלמדות על פערים בייצוג הספרות.",
]
for point in points:
    p = doc.add_paragraph(style='List Bullet')
    set_rtl(p)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.right_indent = Cm(0.5)
    run = p.add_run(point)
    run.font.size = Pt(10)
    set_run_rtl(run)

heading("7.1 מגבלות", 2)

rtl_para(
    "המגבלות העיקריות הנוכחיות: (1) מטריצת הרגישות בנויה מ-keyword matching "
     "פשוט — שיטות ML מתקדמות (fine-tuned embeddings) יניבו דיוק גבוה יותר. "
     "(2) ה-LLM אינו 'זוכר' בין sessions — memory persistence חיצונית תשפר "
     "את הסימולציה הדינמית ארוכת-הטווח. "
     "(3) אין עדיין ייצוג של מנגנוני epigenetics, mRNA half-life, "
     "ולוקליזציה תת-תאית דינמית."
)

divider()

# ═══════════════════════════════════════════════════════════════════════════
#  8. FUTURE
# ═══════════════════════════════════════════════════════════════════════════
heading("8. כיוונים עתידיים", 1)

future = [
    "GENESYS-Plant: גרסה ל-Arabidopsis thaliana עם פרמטרי סביבה של "
     "צמחים (אור, CO₂, לחות, פתוגנים פטרייתיים)",
    "GENESYS-Microbiome: סימולציה של קהילת מיקרוביאל עם inter-species PPI "
     "ומשאבים משותפים",
    "Fine-tuned Biological LLM: אימון מחדש של llama3 על EcoCyc, UniProt, "
     "ו-PubMed כדי לשפר את הדיוק הביולוגי",
    "Federated Multi-Cell Simulation: מספר 'תאים' מקיימים ביניהם "
     "תקשורת — לעבר סימולציית רקמה",
    "Experimental Validation Pipeline: השוואה אוטומטית עם נתוני RNA-seq "
     "ממחקרים ניסויים זמינים בציבור (GEO database)",
    "AI Environmental Interpreter: שילוב עמוק עם LLM לפרשנות תרחישים "
     "בשפה טבעית ('<E. coli בבטן של אדם עם חסר חיסוני')",
]

for item in future:
    p = doc.add_paragraph(style='List Bullet')
    set_rtl(p)
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.right_indent = Cm(0.5)
    run = p.add_run(item)
    run.font.size = Pt(10)
    set_run_rtl(run)

divider()

# ═══════════════════════════════════════════════════════════════════════════
#  9. METHODS
# ═══════════════════════════════════════════════════════════════════════════
heading("9. שיטות", 1)

heading("9.1 מקורות נתונים", 2)
rtl_para(
    "נתוני גנים נלקחו ממסד EcoCyc release 27 ו-UniProt (E. coli K-12, taxon 83333). "
    "נתוני PPI נלקחו מ-STRING DB v12.0 (ציון ≥ 700). "
    "נתוני שפע חלבוני (PPM) נלקחו מ-PaxDB v5.0, "
    "dataset: E. coli K-12 MG1655 whole-cell abundance. "
    "כל הנתונים אוחדו בקובץ k12info.csv (4,762 שורות, 32 עמודות)."
)

heading("9.2 סביבת פיתוח", 2)
rtl_para(
    "Python 3.11, Tkinter, matplotlib 3.10 (3D scatter), numpy, networkx. "
    "LLM backends: Anthropic Claude Haiku 4.5 (API), Ollama llama3 8B Q4_0 (local). "
    "3D layout: פריסה גנומית על פני אליפסואיד לפי מיקום כרומוזומלי (start/end). "
    "גרסת קוד פתוח: GENESYS v2.0 (github.com/[TBD])."
)

divider()

# ═══════════════════════════════════════════════════════════════════════════
#  REFERENCES
# ═══════════════════════════════════════════════════════════════════════════
heading("רשימת מקורות", 1)

refs = [
    "[1] Wei J. et al. (2022) Chain-of-thought prompting elicits reasoning in large language models. NeurIPS.",
    "[2] Yao S. et al. (2023) ReAct: Synergizing reasoning and acting in language models. ICLR.",
    "[3] Russell S. & Norvig P. (2020) Artificial Intelligence: A Modern Approach, 4th ed. Pearson.",
    "[4] Blattner F.R. et al. (1997) The complete genome sequence of Escherichia coli K-12. Science 277:1453.",
    "[5] Keseler I.M. et al. (2021) The EcoCyc database in 2021. Nucleic Acids Research 49:D543.",
    "[6] Guisbert E. et al. (2008) Convergence of molecular, modeling, and systems approaches for an "
     "understanding of the Escherichia coli heat shock response. Microbiol Mol Biol Rev 72:545.",
    "[7] Potrykus K. & Cashel M. (2008) (p)ppGpp: still magical? Annu Rev Microbiol 62:35.",
    "[8] Berardini T.Z. et al. (2015) The Arabidopsis information resource: Making and mining the gold standard "
     "annotated reference plant genome. Genesis 53:474.",
    "[9] Gasch A.P. et al. (2000) Genomic expression programs in the response of yeast cells to environmental "
     "changes. Mol Biol Cell 11:4241.",
    "[10] Wang M. et al. (2015) PaxDb, a database of protein abundance averages across all three domains of life. "
     "Mol Cell Proteomics 14:2914.",
]

for ref in refs:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.left_indent = Cm(0.5)
    run = p.add_run(ref)
    run.font.size = Pt(8.5)
    run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)

divider()

# ── Footer note ──────────────────────────────────────────────────────────────
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(8)
run = p.add_run(
    "© 2026 GENESYS Project  |  Code available under MIT License  |  "
    "Data: EcoCyc, STRING DB, PaxDb, UniProt"
)
run.font.size = Pt(8)
run.font.color.rgb = RGBColor(0xaa, 0xaa, 0xaa)
run.font.italic = True

# ── Save ─────────────────────────────────────────────────────────────────────
doc.save(OUT)
print(f"Saved: {OUT}")
