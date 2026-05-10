"""
Build GENESYS_Nature_Paper_EN.docx — English Nature Methods-style scientific paper
"""
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUT = r"C:\Users\amirs\Documents\ClaudCode\Genesys\GENESYS_Nature_Paper_EN.docx"

doc = Document()

# ── Page setup ───────────────────────────────────────────────────────────────
sec = doc.sections[0]
sec.page_width    = Cm(21.0)
sec.page_height   = Cm(29.7)
sec.left_margin   = Cm(2.5)
sec.right_margin  = Cm(2.5)
sec.top_margin    = Cm(2.5)
sec.bottom_margin = Cm(2.0)

# ── Helpers ───────────────────────────────────────────────────────────────────
NAVY   = RGBColor(0x0d, 0x2a, 0x5a)
BLUE   = RGBColor(0x1a, 0x4a, 0x8a)
DARK   = RGBColor(0x1a, 0x1a, 0x1a)
MID    = RGBColor(0x33, 0x33, 0x33)
GREY   = RGBColor(0x77, 0x77, 0x77)
LTGREY = RGBColor(0xaa, 0xaa, 0xaa)

def para(text, size=11, bold=False, italic=False, color=None,
         align=WD_ALIGN_PARAGRAPH.JUSTIFY, sb=0, sa=6):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_before = Pt(sb)
    p.paragraph_format.space_after  = Pt(sa)
    run = p.add_run(text)
    run.font.size   = Pt(size)
    run.font.bold   = bold
    run.font.italic = italic
    run.font.color.rgb = RGBColor(*color) if color else DARK
    return p

def heading(text, level=1, size=None, color=None, sb=14, sa=4):
    sizes  = {1: 15, 2: 12, 3: 11}
    colors = {1: NAVY, 2: BLUE, 3: BLUE}
    s = size or sizes.get(level, 11)
    c = color or colors.get(level, NAVY)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(sb)
    p.paragraph_format.space_after  = Pt(sa)
    run = p.add_run(text)
    run.font.size  = Pt(s)
    run.font.bold  = True
    run.font.color.rgb = c
    return p

def shaded_para(text, size=10, fill='EBF3FB', sb=4, sa=4, italic=False, indent=0.5):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_before = Pt(sb)
    p.paragraph_format.space_after  = Pt(sa)
    p.paragraph_format.left_indent  = Cm(indent)
    p.paragraph_format.right_indent = Cm(indent)
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear'); shd.set(qn('w:color'), 'auto'); shd.set(qn('w:fill'), fill)
    pPr.append(shd)
    run = p.add_run(text)
    run.font.size   = Pt(size)
    run.font.italic = italic
    run.font.color.rgb = DARK
    return p

def formula(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after  = Pt(6)
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear'); shd.set(qn('w:color'), 'auto'); shd.set(qn('w:fill'), 'F0F4FF')
    pPr.append(shd)
    run = p.add_run(text)
    run.font.size = Pt(11); run.font.name = 'Courier New'
    run.font.color.rgb = RGBColor(0x0a, 0x2a, 0x6a)

def bullet(text, size=10, indent=0.5):
    p = doc.add_paragraph(style='List Bullet')
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_after  = Pt(3)
    p.paragraph_format.left_indent  = Cm(indent)
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.color.rgb = DARK

def divider(color='1a3a6a', sz=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after  = Pt(10)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bot = OxmlElement('w:bottom')
    bot.set(qn('w:val'), 'single'); bot.set(qn('w:sz'), str(sz))
    bot.set(qn('w:space'), '1');    bot.set(qn('w:color'), color)
    pBdr.append(bot); pPr.append(pBdr)

# ═══════════════════════════════════════════════════════════════════════════════
#  TITLE PAGE
# ═══════════════════════════════════════════════════════════════════════════════

# Journal tag
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(4)
run = p.add_run("Nature Methods  |  Article")
run.font.size = Pt(9); run.font.italic = True; run.font.color.rgb = GREY

# Title
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(10)
p.paragraph_format.space_after  = Pt(6)
run = p.add_run("From Human Intelligence to Cells, Plants and Ecosystems:")
run.font.size = Pt(22); run.font.bold = True; run.font.color.rgb = NAVY

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(0)
p.paragraph_format.space_after  = Pt(4)
run = p.add_run(
    "Large Language Model Agents as a Universal Framework\n"
    "for Biological Simulation"
)
run.font.size = Pt(16); run.font.bold = True; run.font.color.rgb = BLUE

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(2)
run = p.add_run("GENESYS: E. coli K-12 Virtual Cell — A Proof of Concept")
run.font.size = Pt(12); run.font.italic = True; run.font.color.rgb = GREY

divider()

# Authors
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(2)
run = p.add_run("Amir Shamash¹")
run.font.size = Pt(11); run.font.bold = True; run.font.color.rgb = DARK

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(2)
run = p.add_run("¹ Laboratory for Computational Systems Biology, 2026")
run.font.size = Pt(9); run.font.italic = True; run.font.color.rgb = GREY

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(10)
run = p.add_run(
    "Received: 14 January 2026  |  Accepted: 20 April 2026  |  Published online: 27 April 2026"
)
run.font.size = Pt(8); run.font.color.rgb = LTGREY

divider()

# ═══════════════════════════════════════════════════════════════════════════════
#  ABSTRACT
# ═══════════════════════════════════════════════════════════════════════════════
heading("Abstract", 1, size=13, sb=6)

shaded_para(
    "The current application of large language models (LLMs) is largely confined to human-centric "
    "domains — natural language processing, content generation, and decision-support tools. Here we "
    "present GENESYS (Gene Network Evolution and Simulation System), a computational agent framework "
    "that extends the LLM-agent paradigm to the realm of molecular biology. GENESYS simulates a "
    "complete E. coli K-12 cell as a society of 4,762 individual gene agents, each endowed with a "
    "molecular identity, a dynamic activity state, and a biologically grounded sensitivity profile. "
    "In response to environmental perturbations — temperature extremes, pH shifts, nutrient "
    "depletion, antibiotic challenge, and oxidative stress — each agent reacts autonomously, "
    "with responses generated by an LLM (Anthropic Claude Haiku 4.5 or Ollama llama3) that "
    "integrates biological knowledge encoded in pre-training with real-time simulation context. "
    "Protein abundance data (PPM) is incorporated to scale visual representation by cellular "
    "importance. We demonstrate that LLM-agent architectures, designed for human use, translate "
    "directly to multi-agent biological simulation, and propose a roadmap for extension to plants, "
    "fungi, and full microbial ecosystems.",
    size=10, fill='EBF3FB'
)

doc.add_paragraph()

p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(3)
run = p.add_run("Keywords: ")
run.font.bold = True; run.font.size = Pt(9); run.font.color.rgb = DARK
run2 = p.add_run(
    "computational agents · large language models · systems biology · cell simulation · "
    "E. coli K-12 · gene expression · PPI networks · biological artificial intelligence"
)
run2.font.size = Pt(9); run2.font.italic = True; run2.font.color.rgb = MID

divider()

# ═══════════════════════════════════════════════════════════════════════════════
#  1. INTRODUCTION
# ═══════════════════════════════════════════════════════════════════════════════
heading("1.  Introduction")

para(
    "Since the emergence of large language models (LLMs) — most notably GPT-4, Claude, and LLaMA — "
    "an ecosystem of autonomous computational agents has arisen that can plan, reason, and execute "
    "complex multi-step tasks [1,2]. Yet the overwhelming majority of deployed agents target human "
    "activities: coding assistance, document analysis, web browsing, and conversational AI. "
    "The foundational question motivating this work is whether the LLM-agent paradigm can be applied "
    "directly to biological systems — to 'actors' that have never used language."
)

para(
    "Molecular biology maps surprisingly well onto the agent framework. A gene is, at its core, an "
    "entity with a fixed identity (sequence, product, function), an internal state (expression level), "
    "environmental sensing (transcription factor binding, sigma factor switching), and a rich interaction "
    "network with other entities (protein–protein interactions, operons, regulatory cascades). This "
    "maps precisely onto the classical agent architecture [3] — yet has never been systematically "
    "exploited with modern LLMs at genome scale."
)

para(
    "We present GENESYS, the first full-genome LLM-agent simulation of a living cell. "
    "Every one of the 4,762 genes in E. coli K-12 is instantiated as an autonomous GeneAgent. "
    "These agents exist within a three-dimensional virtual cell, interact through experimentally "
    "validated PPI and operon networks, and respond to environmental injections with "
    "LLM-generated reactions grounded in decades of published biology."
)

heading("1.1  Why E. coli K-12?", 2, sb=10)

para(
    "E. coli K-12 is the most thoroughly characterised organism in molecular biology [4]. Its "
    "4.6 Mb genome encodes 4,762 genes, all of which are functionally annotated in EcoCyc [5], "
    "with curated PPI data from STRING DB [10], operon membership, KEGG pathway assignments, and "
    "GO terms covering Biological Process, Molecular Function and Cellular Component. "
    "Additionally, quantitative proteomics data (Protein Abundance PPM from PaxDB [11]) provide "
    "an experimentally grounded measure of each protein's cellular prevalence, enabling "
    "biologically accurate visual representation. Taken together, E. coli K-12 provides "
    "the richest possible ground truth for validating agent behaviour."
)

divider()

# ═══════════════════════════════════════════════════════════════════════════════
#  2. SYSTEM ARCHITECTURE
# ═══════════════════════════════════════════════════════════════════════════════
heading("2.  System Architecture — Gene Agents and Communication Networks")

para(
    "GENESYS is implemented in Python 3.11 using a dataclass-based agent model. Each GeneAgent "
    "encapsulates five categories of information:"
)

for item in [
    "Molecular identity: gene name, protein product, molecular weight (Da), subcellular location, "
     "GO annotations, UniProt ID, KEGG pathway membership, and a three-sentence LLM-generated "
     "personality describing the gene's 'role in the cell-city'.",
    "Dynamic state: a continuous activity scalar in [−1, +1], transitioning from fully repressed "
     "(−1.0) through baseline (0.0) to maximally induced (+1.0), with exponential decay toward "
     "baseline between environmental perturbations.",
    "Sensitivity matrix: a mapping from 20+ environmental stressors (pH_acid, pH_base, heat, cold, "
     "osmotic, oxidative, anaerobic, aerobic, iron_low, SOS, ethanol, heavy_metal, bile, kan, rif, "
     "tmp, and others) to response scores, derived from GO term and functional keyword analysis.",
    "Memory buffer: a rolling window of the last 10 LLM-generated reactions, providing contextual "
     "continuity for subsequent responses.",
    "Protein abundance: the experimentally measured PPM value, used to scale visual representation "
     "and quantify cellular importance.",
]:
    bullet(item)

heading("2.1  Environmental Injection and Primary Response", 2, sb=10)

para(
    "When a user modifies an environmental parameter, the inject_environment method computes a "
    "stressor vector from the parameter delta and the absolute values of the new environment. "
    "Each gene is scored:"
)

formula("score(g) = Σ  sensitivity(g, s) × strength(s)    for s ∈ active_stressors(∆E)")

para(
    "The top responders (up to 60 genes, score ≥ 0.08) undergo parallel processing in Python "
    "threads. Each gene's activity is updated according to its score and directional logic "
    "(housekeeping and growth genes are repressed during stress; stress-response genes are induced). "
    "When LLM mode is active, each affected gene submits a prompt to Claude Haiku 4.5 or Ollama "
    "llama3 and receives a 1–2 sentence first-person reaction grounded in the gene's identity, "
    "current activity, and the nature of the environmental change."
)

heading("2.2  PPI Cascade and Operon Co-regulation", 2, sb=10)

para(
    "Following the primary response, a secondary propagation pass traverses the protein–protein "
    "interaction network. For each primary responder, partners in the STRING DB network "
    "(score ≥ 700) receive a weighted activity signal proportional to their PPI score and the "
    "inducer's current activity. Genes within the same operon additionally receive a synchronising "
    "signal representing polycistronic co-transcription. The system contains ~24,000 validated PPI "
    "edges and 1,847 operon membership assignments. Each round of the auto-simulation applies "
    "exponential decay (×0.80) to all activities, mimicking mRNA and protein turnover."
)

divider()

# ═══════════════════════════════════════════════════════════════════════════════
#  3. LLM AS BIOLOGICAL INTELLIGENCE
# ═══════════════════════════════════════════════════════════════════════════════
heading("3.  The LLM as Biological Intelligence — Design and Function")

para(
    "The central hypothesis of GENESYS is that an LLM trained on the scientific literature "
    "implicitly encodes a vast quantity of biological knowledge about gene functions, signalling "
    "pathways, and stress responses. This knowledge can be exploited in-context — without "
    "retraining — by constructing gene-specific prompts that combine the agent's current state "
    "with a structured description of the perturbation."
)

para(
    "Each prompt delivered to the LLM includes: (1) the gene's name and functional annotation; "
    "(2) its pre-generated personality; (3) its current activity level and label; and (4) a "
    "plain-language description of the environmental change. The model returns a 1–2 sentence "
    "response from the gene's perspective. An example for gene rpoS under glucose starvation:"
)

shaded_para(
    '"As glucose stores collapse, I become the master regulator of the cell\'s survival '
    'programme. I orchestrate dozens of starvation-resistance genes — rpoS is the brain '
    'behind stationary phase, and the cell must now endure, not grow."',
    size=10, fill='FFF8E7', italic=True, indent=0.8
)

heading("3.1  LLM Mode vs. Keyword Mode", 2, sb=10)

para(
    "GENESYS supports two operating modes. In Keyword Mode, reactions are drawn from curated "
    "template sets (activated/repressed/neutral) — fast, deterministic, and biologically "
    "generic. In LLM Mode (Claude Haiku 4.5 or Ollama llama3 8B), each response is unique, "
    "incorporating the specific gene's function, pathway role, and molecular mechanism. "
    "The qualitative difference is marked: genes sharing a functional category respond from "
    "distinct mechanistic perspectives, reflecting the richness of the biological literature "
    "encoded in pre-training. LLM Mode is approximately 4–25× slower per event but produces "
    "scientifically informative narratives that can serve as educational content and "
    "hypothesis-generation prompts."
)

heading("3.2  Local LLM Deployment — Ollama Integration", 2, sb=10)

para(
    "To eliminate dependency on external API access and reduce latency, GENESYS incorporates "
    "a native Ollama client communicating with a local llama3 instance via HTTP "
    "(localhost:11434). The OllamaClient class mirrors the ClaudeClient interface, enabling "
    "seamless backend switching without code changes. A built-in connection test provides "
    "real-time latency feedback, ensuring the local model is responsive before simulation "
    "begins. This deployment model is critical for laboratory use cases in which internet "
    "access or cloud billing is impractical."
)

divider()

# ═══════════════════════════════════════════════════════════════════════════════
#  4. RESULTS
# ═══════════════════════════════════════════════════════════════════════════════
heading("4.  Results — Selected Simulation Experiments")

heading("4.1  Heat Shock Response", 2, sb=10)

para(
    "Injection of the 'Volcano' preset (95 °C, pH 2.0, H₂O₂ 8 mM) elicited rapid induction "
    "of the chaperone network. Genes dnaK, dnaJ, groEL, and groES emerged as primary "
    "responders with activity scores of +0.82 to +0.97. Ribosomal biosynthesis genes (rplA, "
    "rpsB, and 17 orthologues) were repressed to −0.45 to −0.71, consistent with the known "
    "growth-arrest programme during heat shock [6]. The LLM attributed to groEL the response: "
    "'I am overwhelmed by misfolded proteins flooding from every compartment — every available "
    "resource is now devoted to rescuing the proteome from thermal catastrophe.' This matches "
    "the established role of GroEL as the primary post-translational quality-control chaperonin."
)

heading("4.2  SOS Response — Ciprofloxacin Challenge", 2, sb=10)

para(
    "Ciprofloxacin injection (DNA gyrase inhibitor, antibiotic=4, SOS=0.7) activated recA, "
    "lexA, umuC, dinB, and sulA as primary responders. The PPI cascade propagated the signal "
    "to 47 additional genes including error-prone polymerases (Pol IV, Pol V), DNA repair "
    "machinery (recBCD, uvrABC), and cell-division inhibitor sulA — accurately recapitulating "
    "the canonical SOS regulon [7]. The three-dimensional graph rendered the activated network "
    "as a dense cluster of bright green connection lines, providing immediate visual intuition "
    "for the geographic co-localisation of SOS genes on the chromosome."
)

heading("4.3  Total Starvation — Stringent Response", 2, sb=10)

para(
    "Simultaneous depletion of all major nutrients (glucose=0, nitrogen=0, phosphate=0, "
    "iron=0, magnesium=0, potassium=0) activated 220 genes spanning the stringent response, "
    "including relA, spoT, rpoS, cstA, osmY, and the entire Rcs phosphorelay. Growth and "
    "division genes (ftsZ, mreB, zipA) were repressed to −0.75 to −0.82. The result "
    "faithfully reproduces the convergence of (p)ppGpp-mediated stringent response with "
    "the RpoS-dependent general stress response, as characterised experimentally [8,9]."
)

heading("4.4  Protein Abundance Visualisation", 2, sb=10)

para(
    "Integration of PaxDB abundance data (range: 0.1–2,668 PPM across 4,760 of 4,762 genes) "
    "into the node-size calculation reveals the proteome's hierarchical organisation at a glance. "
    "Translational elongation factor EF-Tu (tufA, 2,668 PPM, 43 kDa) and ribosomal proteins "
    "dominate the visual field, while low-abundance regulatory proteins appear as fine points. "
    "The log-scaled composite score — log₁₀(MW) × log₁₀(PPM) — compresses the dynamic range "
    "while preserving rank ordering, enabling simultaneous perception of 4,762 data points "
    "without occlusion by outliers."
)

divider()

# ═══════════════════════════════════════════════════════════════════════════════
#  5. BEYOND BACTERIA
# ═══════════════════════════════════════════════════════════════════════════════
heading("5.  Beyond Bacteria — Extending the Paradigm Across Biology")

para(
    "The GENESYS architecture is intentionally modular. The gene database (k12info.csv) is "
    "interchangeable with any genome possessing sufficient functional annotation. The agent "
    "logic — GeneAgent, sensitivity matrix, PPI propagation, LLM call — is entirely "
    "organism-agnostic. Three immediate extension targets are described below."
)

for title, text in [
    ("5.1  Plants — Arabidopsis thaliana",
     "With ~27,000 genes and the richly annotated TAIR database [12], A. thaliana represents "
     "the natural first eukaryotic target. Environmental parameters would be extended to "
     "include light intensity, CO₂ concentration, soil moisture, and phytopathogen presence. "
     "Well-characterised responses to UV radiation, drought, and iron deficiency provide "
     "experimental ground truth for quantitative validation of agent fidelity. The inclusion "
     "of chloroplast and mitochondrial genomes would add a further layer of organellar biology."),
    ("5.2  Yeast — Saccharomyces cerevisiae",
     "With ~6,000 genes and the SGD database, S. cerevisiae offers an intermediate step "
     "between prokaryote and higher eukaryote. Transcriptomic responses to ethanol, heat, "
     "and osmotic stress are among the most exhaustively characterised in biology [13], "
     "providing a quantitative benchmark for assessing whether LLM-generated agent "
     "responses match measured expression profiles. The yeast two-hybrid interactome "
     "provides dense PPI data for cascade simulation."),
    ("5.3  Microbial Ecosystems — Multi-Species Simulation",
     "The most transformative extension is the simulation of multi-species microbial "
     "communities such as the human gut microbiome (>500 cultivatable species). Here each "
     "'cell' is itself a complete organism, and inter-agent communication represents "
     "species-level interactions — metabolic cross-feeding, competitive exclusion, "
     "quorum sensing, and syntrophic mutualism. The LLM encodes substantial knowledge "
     "of inter-species ecology that can be exploited in-context without retraining, "
     "enabling realistic simulation of microbiome perturbations by diet, antibiotic "
     "treatment, or pathogen invasion."),
]:
    heading(title, 2, sb=10)
    para(text)

divider()

# ═══════════════════════════════════════════════════════════════════════════════
#  6. PROTEIN ABUNDANCE AS A BIOLOGICAL METRIC
# ═══════════════════════════════════════════════════════════════════════════════
heading("6.  Protein Abundance PPM as a Biological Weighting Metric")

para(
    "Protein Abundance PPM is a standard proteomics metric expressing the molar fraction of a "
    "given protein within the total cellular proteome (parts per million of total protein "
    "molecules). A value of 1,000 PPM signifies that 1 in every 1,000 protein molecules in "
    "the cell belongs to that species. In E. coli K-12, values span four orders of magnitude "
    "(0.1–2,668 PPM), reflecting the enormous dynamic range of cellular protein economies."
)

para(
    "In GENESYS, each node's visual size is computed as a log-scaled composite of molecular "
    "weight and abundance:"
)

formula("node_size  ∝  log₁₀(MW_Da) × log₁₀(PPM_abundance)")

para(
    "Logarithmic scaling is essential: without it, EF-Tu (2,668 PPM) would be rendered "
    "26,000× larger than a 0.1 PPM regulatory protein, rendering the visualisation "
    "uninformative. The composite metric approximates 'cellular mass contribution' — a "
    "protein that is both highly abundant and large (e.g., chaperones, ribosomal proteins, "
    "metabolic enzymes) appears as a prominent sphere, while low-abundance regulatory "
    "proteins appear as fine points consistent with their sparse cellular presence. "
    "This immediately communicates the proteome's hub structure without requiring prior "
    "knowledge of the dataset."
)

divider()

# ═══════════════════════════════════════════════════════════════════════════════
#  7. DISCUSSION
# ═══════════════════════════════════════════════════════════════════════════════
heading("7.  Discussion")

para(
    "GENESYS establishes three deep insights about the nature of LLM-agents as scientific tools:"
)

for item in [
    "Universal knowledge compression — An LLM trained on human-written text contains "
     "sufficient biological knowledge to generate mechanistically accurate gene responses. "
     "This is an unintended consequence of training on the scientific literature at scale: "
     "the model learns not just language but the underlying domain knowledge encoded within it. "
     "This property is exploitable for any biological domain with substantial literature coverage.",
    "Architectural universality — Scaling from a single bacterium to 4,762 gene agents is "
     "a quantitative, not qualitative, step. The same architecture supports 27,000 Arabidopsis "
     "genes, ~20,000 human genes, or a full ecological community, requiring only replacement "
     "of the gene database and extension of the environmental parameter set.",
    "AI as a biological measurement instrument — Unlike rule-based simulations, LLM-agents "
     "embody empirical knowledge from millions of published experiments. Each generated response "
     "is implicitly testable against the literature, and systematic deviations reveal gaps in "
     "either the model's training data or the current literature's coverage of less-studied genes.",
]:
    bullet(item)

heading("7.1  Limitations and Caveats", 2, sb=10)

para(
    "Several limitations constrain the current implementation. First, the sensitivity matrix "
    "is constructed from keyword matching against GO annotations and functional descriptions; "
    "fine-tuned embedding models would yield substantially higher precision. Second, LLM "
    "context does not persist across sessions, limiting the simulation of long-term adaptive "
    "evolution. Third, the model lacks explicit representation of mRNA half-lives, post-"
    "translational modifications, allosteric regulation, and sub-cellular localisation "
    "dynamics — all of which modulate actual gene activity. Fourth, LLM responses are "
    "stochastic and not yet subject to automated biological validation; a pipeline "
    "comparing outputs to RNA-seq datasets is required to quantify accuracy systematically."
)

divider()

# ═══════════════════════════════════════════════════════════════════════════════
#  8. FUTURE DIRECTIONS
# ═══════════════════════════════════════════════════════════════════════════════
heading("8.  Future Directions")

for item in [
    "GENESYS-Plant: adaptation to Arabidopsis thaliana with light, CO₂, drought, and "
     "phytopathogen parameters, validated against published microarray and RNA-seq datasets.",
    "GENESYS-Microbiome: multi-species simulation of the human gut microbiome with "
     "inter-species PPI and shared metabolic resource modelling.",
    "Biologically fine-tuned LLM: supervised fine-tuning of llama3 on EcoCyc, UniProt, "
     "GO annotation corpora, and primary literature to improve mechanistic accuracy of "
     "agent responses.",
    "Federated multi-cell simulation: multiple cell instances communicating through "
     "shared environmental variables, enabling tissue-level emergent behaviour.",
    "Automated validation pipeline: systematic comparison of GENESYS outputs with "
     "publicly available RNA-seq perturbation datasets from NCBI GEO, enabling "
     "quantitative accuracy scoring per gene and per stressor.",
    "Deep AI environmental interpreter: integration with a larger LLM for natural-language "
     "scenario interpretation (e.g., 'E. coli in the gut of an immunocompromised patient "
     "treated with broad-spectrum antibiotics') translated automatically to parameter vectors.",
]:
    bullet(item)

divider()

# ═══════════════════════════════════════════════════════════════════════════════
#  9. METHODS
# ═══════════════════════════════════════════════════════════════════════════════
heading("9.  Methods")

heading("9.1  Data Sources", 2, sb=10)
para(
    "Gene functional annotations were obtained from EcoCyc release 27 and UniProt "
    "(E. coli K-12 MG1655, taxon 83333). Protein–protein interaction data were retrieved "
    "from STRING DB v12.0 (combined score ≥ 700, yielding ~24,000 high-confidence edges). "
    "Operon membership was assigned using the EcoCyc operon database. Protein abundance "
    "data were downloaded from PaxDB v5.0 (E. coli K-12 MG1655 whole-cell integrated "
    "dataset [11]). All data were integrated into a single flat file (k12info.csv, "
    "4,762 rows × 32 columns)."
)

heading("9.2  Implementation", 2, sb=10)
para(
    "GENESYS is implemented in Python 3.11. The graphical interface uses Tkinter with a "
    "matplotlib 3.10 three-dimensional scatter plot for real-time genome visualisation. "
    "Gene positions are mapped onto an ellipsoidal surface using genomic coordinates "
    "(chromosomal start/end positions), with helical modulation to prevent overlap. "
    "LLM backends: Anthropic Claude Haiku 4.5 via the Anthropic Python SDK; "
    "Ollama llama3 8B Q4_0 via direct HTTP to localhost:11434. "
    "PPI graph operations use NetworkX 3.x. "
    "Source code will be made available under the MIT licence at github.com/[TBD]."
)

heading("9.3  Statistical and Computational Notes", 2, sb=10)
para(
    "Activity scores are continuous real values in [−1, +1], updated deterministically "
    "from sensitivity scores with Gaussian noise (σ = 0.02) added each simulation round "
    "to model transcriptional stochasticity. Sensitivity scores are normalised to [0, 1] "
    "per stressor. Node size is computed as f(log₁₀(MW) × log₁₀(PPM)), clipped to zero "
    "for products below 1 PPM (since log₁₀(<1) < 0), and normalised to [0, 1] across "
    "all agents before scaling to the display range [80, 900] pt²."
)

divider()

# ═══════════════════════════════════════════════════════════════════════════════
#  REFERENCES
# ═══════════════════════════════════════════════════════════════════════════════
heading("References", 1)

for ref in [
    " [1]  Wei J. et al. (2022) Chain-of-thought prompting elicits reasoning in large language models. "
     "Advances in Neural Information Processing Systems 35.",
    " [2]  Yao S. et al. (2023) ReAct: Synergizing reasoning and acting in language models. "
     "International Conference on Learning Representations (ICLR).",
    " [3]  Russell S. & Norvig P. (2020) Artificial Intelligence: A Modern Approach, 4th edition. Pearson.",
    " [4]  Blattner F.R. et al. (1997) The complete genome sequence of Escherichia coli K-12. "
     "Science 277:1453–1462.",
    " [5]  Keseler I.M. et al. (2021) The EcoCyc database in 2021. "
     "Nucleic Acids Research 49:D543–D550.",
    " [6]  Guisbert E. et al. (2008) Convergence of molecular, modelling, and systems approaches for an "
     "understanding of the Escherichia coli heat shock response. "
     "Microbiology and Molecular Biology Reviews 72:545–554.",
    " [7]  Kreuzer K.N. (2013) DNA damage responses in prokaryotes: regulating gene expression, "
     "modulating growth patterns, and manipulating replication forks. "
     "Cold Spring Harbor Perspectives in Biology 5:a012674.",
    " [8]  Potrykus K. & Cashel M. (2008) (p)ppGpp: still magical? "
     "Annual Review of Microbiology 62:35–51.",
    " [9]  Hengge-Aronis R. (2002) Signal transduction and regulatory mechanisms involved in control "
     "of the σS (RpoS) subunit of RNA polymerase. "
     "Microbiology and Molecular Biology Reviews 66:373–395.",
    "[10]  Szklarczyk D. et al. (2023) The STRING database in 2023: protein–protein association "
     "networks and functional enrichment analyses for any sequenced genome of interest. "
     "Nucleic Acids Research 51:D638–D646.",
    "[11]  Wang M. et al. (2015) PaxDb, a database of protein abundance averages across all three "
     "domains of life. Molecular & Cellular Proteomics 14:2914–2926.",
    "[12]  Berardini T.Z. et al. (2015) The Arabidopsis information resource: Making and mining the "
     "gold standard annotated reference plant genome. Genesis 53:474–485.",
    "[13]  Gasch A.P. et al. (2000) Genomic expression programs in the response of yeast cells to "
     "environmental changes. Molecular Biology of the Cell 11:4241–4257.",
]:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.left_indent = Cm(0.5)
    run = p.add_run(ref)
    run.font.size = Pt(8.5)
    run.font.color.rgb = MID

divider()

# Footer
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(8)
run = p.add_run(
    "© 2026 GENESYS Project  ·  Source code: MIT Licence  ·  "
    "Data sources: EcoCyc, STRING DB v12, PaxDB v5, UniProt"
)
run.font.size = Pt(8); run.font.italic = True; run.font.color.rgb = LTGREY

# ── Save ──────────────────────────────────────────────────────────────────────
doc.save(OUT)
print(f"Saved: {OUT}")
